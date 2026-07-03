# -*- coding: utf-8 -*-
"""
勘察风险分析报告生成引擎
========================
支持三种产出格式，全部本地生成，无需外部服务：
  - Markdown (.md)   —— 轻量、纯文本、可版本控制
  - Word (.docx)     —— 正式工程交付文档，含封面/表格/嵌入图（python-docx）
  - HTML (.html)     —— 浏览器可直接打开/打印为 PDF，含样式与嵌入图

支持两种报告范围：
  - 单风险报告  generate_risk_report(rid, fmt)
  - 全线综合报告 generate_full_report(fmt)

设计原则：
  1) 内容由证据链驱动 —— 每个结论可追溯至影像/点云/物探/钻孔/报告
  2) 工程结构完整 —— 封面、概述、分项、参数表、地层表、嵌入图、结论建议
  3) 多格式同源 —— 三种格式内容一致，只是渲染方式不同
"""
import os
import io
import json
import datetime
from typing import Optional, List, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(HERE, "data")

# ----------------------------------------------------------------------------
# 加载数据
# ----------------------------------------------------------------------------
def _load(*path):
    with open(os.path.join(DATA, *path), "r", encoding="utf-8") as f:
        return json.load(f)

MANIFEST = _load("manifest.json")
RISKS = MANIFEST["risk_objects"]
RISK_BY_ID = {r["id"]: r for r in RISKS}
BOREHOLES = _load("boreholes", "boreholes.json")
BH_BY_ID = {b["id"]: b for b in BOREHOLES}
GEO_LINES = _load("geophysics", "lines.json")
GEO_BY_ID = {g["id"]: g for g in GEO_LINES}
REPORT = _load("report", "survey_report.json")

LEVEL_RANK = {"高": 3, "中高": 2, "中": 1}


def _img_path(rel: str) -> str:
    return os.path.join(DATA, rel.replace("/", os.sep))


def _now() -> str:
    return datetime.datetime.now().strftime("%Y-%m-%d")


# #############################################################################
#  A. 数据收集层 —— 把分散的数据组织成报告需要的结构
# #############################################################################
def collect_risk_context(rid: str) -> Dict[str, Any]:
    """收集单个风险的完整上下文（含关联钻孔、物探、报告段落）。"""
    r = RISK_BY_ID[rid]
    bhs = [BH_BY_ID[bid] for bid in r.get("borehole_ids", []) if bid in BH_BY_ID]
    geo = GEO_BY_ID.get(r.get("geophysics_line"))
    related = [s for s in REPORT["sections"] if rid in s.get("related_risks", [])]
    return {"risk": r, "boreholes": bhs, "geophysics": geo, "report_sections": related}


def collect_full_context() -> Dict[str, Any]:
    """收集全线综合报告的上下文。"""
    sections = [collect_risk_context(r["id"]) for r in RISKS]
    # 按里程排序
    sections.sort(key=lambda s: s["risk"]["mileage_m"])
    return {
        "project": MANIFEST["project"],
        "route": MANIFEST["route"],
        "risks": RISKS,
        "boreholes": BOREHOLES,
        "geo_lines": GEO_LINES,
        "report": REPORT,
        "risk_sections": sections,
    }


# #############################################################################
#  B. Markdown 生成器
# #############################################################################
def _md_risk_section(ctx: Dict, level: int = 2) -> str:
    """生成单个风险的 Markdown 段落。"""
    r, bhs, geo, related = ctx["risk"], ctx["boreholes"], ctx["geophysics"], ctx["report_sections"]
    h = "#" * level
    e = r["evidence"]
    md = [f"{h} {r['name']}\n"]
    md.append(f"- **里程**：{r['mileage']}")
    md.append(f"- **风险类型**：{r['type_cn']}")
    md.append(f"- **风险等级**：**{r['risk_level']}**")
    md.append(f"- **评估可信度**：{r['confidence']}")
    md.append(f"- **关联钻孔**：{', '.join(b['id'] for b in bhs) or '—'}")
    md.append(f"- **关联物探**：{geo['name'] if geo else '—'}\n")

    md.append(f"{h}# 多源证据\n")
    md.append("| 数据源 | 证据描述 |")
    md.append("|---|---|")
    md.append(f"| 🛰 正射影像 | {e['image']} |")
    md.append(f"| 🏔 三维点云 | {e['pointcloud']} |")
    md.append(f"| 📡 物探剖面 | {e['geophysics']} |")
    md.append(f"| 🔩 钻孔资料 | {e['borehole']} |")
    md.append(f"| 📄 勘察报告 | {e['report']} |\n")

    if e.get("params"):
        md.append(f"{h}# 关键参数\n")
        md.append("| 参数 | 数值 |")
        md.append("|---|---|")
        for k, v in e["params"].items():
            md.append(f"| {_param_label(k)} | {v} |")
        md.append("")

    md.append(f"{h}# 综合风险解释\n")
    md.append(r["interpretation"] + "\n")

    md.append(f"{h}# 设计与施工建议\n")
    md.append(r["design_suggestion"] + "\n")

    if bhs:
        md.append(f"{h}# 钻孔地层资料（验证依据）\n")
        for b in bhs:
            md.append(f"**{b['id']}**（{b['mileage']}，孔口高程 {b['elevation']}m，"
                      f"孔深 {b['depth_m']}m" +
                      (f"，地下水位埋深 {b['water_depth_m']}m" if b.get("water_depth_m") is not None else "")
                      + "）\n")
            md.append("| 起(m) | 止(m) | 岩性 | 描述 |")
            md.append("|---|---|---|---|")
            for L in b["layers"]:
                md.append(f"| {L['top']} | {L['bottom']} | {L['lithology']} | {L['desc']} |")
            md.append("")
    return "\n".join(md)


def _param_label(k):
    return {"max_slope_deg": "最大坡度(°)", "relief_m": "相对高差(m)",
            "rho_min": "最低电阻率(Ω·m)", "weathered_depth_m": "风化层厚度(m)",
            "water_depth_m": "地下水位埋深(m)", "fracture_width_m": "破碎带宽度(m)",
            "rqd_pct": "RQD(%)", "avg_slope_deg": "平均坡度(°)",
            "deposit_depth_m": "堆积层厚度(m)"}.get(k, k)


def generate_risk_markdown(rid: str) -> str:
    ctx = collect_risk_context(rid)
    r = ctx["risk"]
    md = []
    md.append(f"# {r['name']} 风险分析报告\n")
    md.append(f"> **项目**：{MANIFEST['project']['scenario']}  ")
    md.append(f"> **生成时间**：{_now()}  ")
    md.append(f"> **生成方式**：多源勘察数据联动展示与证据链追溯系统（自动生成）\n")
    md.append("---\n")
    md.append(_md_risk_section(ctx, level=1))
    md.append("\n---\n")
    md.append("*本报告由系统基于多源勘察证据链自动生成，所有结论均可追溯至"
              "影像、点云、物探、钻孔、勘察报告等具体数据来源。*\n")
    return "\n".join(md)


def generate_full_markdown() -> str:
    ctx = collect_full_context()
    md = []
    # 封面信息
    md.append(f"# {MANIFEST['project']['scenario']}\n")
    md.append(f"## 多源勘察风险分析综合报告\n")
    md.append(f"> **报告类型**：全线综合风险分析  ")
    md.append(f"> **里程范围**：{MANIFEST['route']['start_mileage']} ~ {MANIFEST['route']['end_mileage']}  ")
    md.append(f"> **生成时间**：{_now()}  ")
    md.append(f"> **生成方式**：多源勘察数据联动展示与证据链追溯系统（自动生成）\n")
    md.append("---\n")

    # 1. 工程概述
    md.append("## 一、工程概述\n")
    overview = next((s for s in REPORT["sections"] if s["id"] == "1"), None)
    if overview:
        md.append(overview["content"] + "\n")
    md.append(f"- **线路长度**：1000m（{MANIFEST['route']['start_mileage']} ~ {MANIFEST['route']['end_mileage']}）")
    md.append(f"- **钻孔数量**：{len(BOREHOLES)} 个")
    md.append(f"- **物探测线**：{len(GEO_LINES)} 条（高密度电法）")
    md.append(f"- **识别风险**：{len(RISKS)} 个"
              f"（高 {sum(1 for r in RISKS if r['risk_level']=='高')}，"
              f"中高 {sum(1 for r in RISKS if r['risk_level']=='中高')}，"
              f"中 {sum(1 for r in RISKS if r['risk_level']=='中')}）\n")

    # 2. 风险统计
    md.append("## 二、风险统计\n")
    md.append("| 风险编号 | 里程 | 风险类型 | 风险等级 | 可信度 |")
    md.append("|---|---|---|---|---|")
    for r in sorted(RISKS, key=lambda x: x["mileage_m"]):
        md.append(f"| {r['id']} | {r['mileage']} | {r['type_cn']} | **{r['risk_level']}** | {r['confidence']} |")
    md.append("")

    # 3. 逐风险分析
    md.append("## 三、逐风险多源证据分析\n")
    for sec in ctx["risk_sections"]:
        md.append(_md_risk_section(sec, level=3))
        md.append("---\n")

    # 4. 风险对比
    if len(RISKS) >= 2:
        md.append("## 四、风险对比分析\n")
        sr = sorted(RISKS, key=lambda x: -LEVEL_RANK.get(x["risk_level"], 0))
        md.append("| 对比项 | " + " | ".join(f"{r['mileage']}" for r in sr[:3]) + " |")
        md.append("|---|" + "---|" * len(sr[:3]))
        md.append("| 风险类型 | " + " | ".join(r["type_cn"] for r in sr[:3]) + " |")
        md.append("| 风险等级 | " + " | ".join(f"**{r['risk_level']}**" for r in sr[:3]) + " |")
        md.append("| 最大坡度(°) | " + " | ".join(
            str(r["evidence"]["params"].get("max_slope_deg", r["evidence"]["params"].get("avg_slope_deg", "—")))
            for r in sr[:3]) + " |")
        md.append("| 最低电阻率(Ω·m) | " + " | ".join(
            str(r["evidence"]["params"].get("rho_min", "—")) for r in sr[:3]) + " |")
        md.append("")
        top = sr[0]
        md.append(f"**首要风险**：{top['mileage']} {top['type_cn']}（{top['risk_level']}），"
                  f"需优先处置。\n")

    # 5. 总体结论建议
    md.append("## 五、总体结论与建议\n")
    concl = next((s for s in REPORT["sections"] if s["id"] == "8"), None)
    if concl:
        md.append(concl["content"] + "\n")
    md.append("### 综合建议\n")
    md.append("1. **按风险等级分优先级处置**：高风险区（K12+380 边坡）应在进洞前完成处治；")
    md.append("2. **超前地质预报全过程跟进**：特别是富水破碎带段，采用 TSP + 地质雷达 + 超前钻孔；")
    md.append("3. **监控量测体系**：建立坡体位移、洞内收敛、地下水动态监测，动态调整支护；")
    md.append("4. **应急准备**：富水段备用抽排水能力 ≥200m³/h，设置防水闸门。\n")

    md.append("---\n")
    md.append("*本报告由多源勘察数据联动展示与证据链追溯系统基于正射影像、三维点云、"
              "物探剖面、钻孔资料、勘察报告等多源证据链自动生成。"
              "所有结论均可追溯至具体数据来源。*\n")
    return "\n".join(md)


# #############################################################################
#  C. Word (.docx) 生成器 —— 正式工程交付文档
# #############################################################################
def _set_cell_shading(cell, color_hex):
    """给表格单元格设置背景色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _add_table(doc, headers: List[str], rows: List[List[str]], header_color="1F4E79"):
    """添加带样式的表格。"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, header_color)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    return t


def _new_doc() -> Document:
    doc = Document()
    # 默认中文字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return doc


def _add_heading_cn(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h


def _docx_risk_section(doc, ctx: Dict):
    """在 Word 文档中追加单个风险分析章节。"""
    r, bhs, geo, related = ctx["risk"], ctx["boreholes"], ctx["geophysics"], ctx["report_sections"]
    e = r["evidence"]
    _add_heading_cn(doc, f"{r['name']}", level=2)

    # 基本信息
    _add_heading_cn(doc, "基本信息", level=3)
    _add_table(doc,
        ["项目", "内容"],
        [["里程", r["mileage"]], ["风险类型", r["type_cn"]],
         ["风险等级", r["risk_level"]], ["评估可信度", r["confidence"]],
         ["关联钻孔", ", ".join(b["id"] for b in bhs) or "—"],
         ["关联物探", geo["name"] if geo else "—"]])

    # 多源证据表
    _add_heading_cn(doc, "多源证据", level=3)
    _add_table(doc,
        ["数据源", "证据描述"],
        [["正射影像", e["image"]], ["三维点云", e["pointcloud"]],
         ["物探剖面", e["geophysics"]], ["钻孔资料", e["borehole"]],
         ["勘察报告", e["report"]]])

    # 关键参数
    if e.get("params"):
        _add_heading_cn(doc, "关键参数", level=3)
        _add_table(doc, ["参数", "数值"],
                   [[_param_label(k), str(v)] for k, v in e["params"].items()])

    # 嵌入物探剖面图
    if geo and os.path.exists(_img_path(geo["image"])):
        _add_heading_cn(doc, "物探剖面图", level=3)
        doc.add_picture(_img_path(geo["image"]), width=Cm(15))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 嵌入钻孔柱状图
    if bhs:
        _add_heading_cn(doc, "钻孔柱状图", level=3)
        for b in bhs:
            bp = _img_path(f"boreholes/{b['id']}.png")
            if os.path.exists(bp):
                doc.add_picture(bp, width=Cm(8))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 综合解释
    _add_heading_cn(doc, "综合风险解释", level=3)
    p = doc.add_paragraph(r["interpretation"])
    p.paragraph_format.first_line_indent = Cm(0.8)

    # 设计建议
    _add_heading_cn(doc, "设计与施工建议", level=3)
    p = doc.add_paragraph(r["design_suggestion"])
    p.paragraph_format.first_line_indent = Cm(0.8)

    # 钻孔地层表
    if bhs:
        _add_heading_cn(doc, "钻孔地层资料", level=3)
        for b in bhs:
            doc.add_paragraph(
                f"{b['id']}（{b['mileage']}，孔口高程 {b['elevation']}m，孔深 {b['depth_m']}m"
                + (f"，地下水位埋深 {b['water_depth_m']}m" if b.get("water_depth_m") is not None else "")
                + "）").runs[0].bold = True
            _add_table(doc, ["起(m)", "止(m)", "岩性", "描述"],
                       [[str(L["top"]), str(L["bottom"]), L["lithology"], L["desc"]]
                        for L in b["layers"]])
            doc.add_paragraph()


def generate_risk_docx(rid: str) -> bytes:
    """生成单风险 Word 报告，返回字节流。"""
    ctx = collect_risk_context(rid)
    r = ctx["risk"]
    doc = _new_doc()
    # 封面标题
    title = doc.add_heading("", level=0)
    run = title.add_run(f"{r['name']}\n风险分析报告")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rr in title.runs:
        rr.font.name = "微软雅黑"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"\n项目：{MANIFEST['project']['scenario']}\n"
                 f"生成时间：{_now()}\n"
                 f"生成方式：多源勘察数据联动展示与证据链追溯系统（自动生成）\n").font.size = Pt(10)
    doc.add_page_break()

    _docx_risk_section(doc, ctx)
    doc.add_paragraph()
    p = doc.add_paragraph("—— 本报告由系统基于多源勘察证据链自动生成，"
                          "所有结论均可追溯至影像、点云、物探、钻孔、勘察报告等具体数据来源。 ——")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_full_docx() -> bytes:
    """生成全线综合 Word 报告，返回字节流。"""
    ctx = collect_full_context()
    doc = _new_doc()
    # 封面
    title = doc.add_heading("", level=0)
    run = title.add_run(f"{MANIFEST['project']['scenario']}\n多源勘察风险分析综合报告")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rr in title.runs:
        rr.font.name = "微软雅黑"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"\n报告类型：全线综合风险分析\n"
                 f"里程范围：{MANIFEST['route']['start_mileage']} ~ {MANIFEST['route']['end_mileage']}\n"
                 f"生成时间：{_now()}\n"
                 f"生成方式：多源勘察数据联动展示与证据链追溯系统（自动生成）\n").font.size = Pt(10)
    doc.add_page_break()

    # 一、工程概述
    _add_heading_cn(doc, "一、工程概述", level=1)
    overview = next((s for s in REPORT["sections"] if s["id"] == "1"), None)
    if overview:
        doc.add_paragraph(overview["content"])
    _add_table(doc, ["项目", "数据"],
        [["线路长度", "1000m"],
         ["钻孔数量", f"{len(BOREHOLES)} 个"],
         ["物探测线", f"{len(GEO_LINES)} 条（高密度电法）"],
         ["识别风险", f"{len(RISKS)} 个"]])

    # 二、风险统计
    _add_heading_cn(doc, "二、风险统计", level=1)
    sr = sorted(RISKS, key=lambda x: x["mileage_m"])
    _add_table(doc, ["风险编号", "里程", "风险类型", "风险等级", "可信度"],
               [[r["id"], r["mileage"], r["type_cn"], r["risk_level"], r["confidence"]] for r in sr])

    # 三、逐风险分析
    _add_heading_cn(doc, "三、逐风险多源证据分析", level=1)
    for sec in ctx["risk_sections"]:
        _docx_risk_section(doc, sec)
        doc.add_paragraph()

    # 四、总体结论建议
    _add_heading_cn(doc, "四、总体结论与建议", level=1)
    concl = next((s for s in REPORT["sections"] if s["id"] == "8"), None)
    if concl:
        doc.add_paragraph(concl["content"])
    _add_heading_cn(doc, "综合建议", level=2)
    for s in ["按风险等级分优先级处置：高风险区（K12+380 边坡）应在进洞前完成处治。",
              "超前地质预报全过程跟进：特别是富水破碎带段，采用 TSP + 地质雷达 + 超前钻孔。",
              "监控量测体系：建立坡体位移、洞内收敛、地下水动态监测，动态调整支护。",
              "应急准备：富水段备用抽排水能力 ≥200m³/h，设置防水闸门。"]:
        doc.add_paragraph(s, style="List Number")

    doc.add_paragraph()
    p = doc.add_paragraph("—— 本报告由多源勘察数据联动展示与证据链追溯系统基于正射影像、"
                          "三维点云、物探剖面、钻孔资料、勘察报告等多源证据链自动生成。"
                          "所有结论均可追溯至具体数据来源。 ——")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# #############################################################################
#  D. HTML 生成器（含样式，可浏览器打印为 PDF）
# #############################################################################
HTML_CSS = """
<style>
  body{font-family:'Microsoft YaHei','SimSun',sans-serif;max-width:900px;margin:30px auto;
       padding:0 30px;color:#222;line-height:1.75;font-size:14px}
  h1{color:#1F4E79;border-bottom:3px solid #1F4E79;padding-bottom:10px;font-size:22px}
  h2{color:#2E75B6;border-left:5px solid #2E75B6;padding-left:10px;margin-top:30px;font-size:18px}
  h3{color:#333;margin-top:20px;font-size:15px}
  table{border-collapse:collapse;width:100%;margin:12px 0;font-size:13px}
  th{background:#1F4E79;color:#fff;padding:8px 10px;text-align:left}
  td{border:1px solid #ccc;padding:6px 10px}
  tr:nth-child(even){background:#f5f8fc}
  img{max-width:100%;border:1px solid #ddd;margin:8px 0;border-radius:4px}
  blockquote{border-left:4px solid #2E75B6;margin:10px 0;padding:8px 16px;
             background:#f0f6fc;color:#555;font-size:13px}
  .cover{text-align:center;padding:60px 0}
  .cover h1{border:none;font-size:28px;margin-bottom:20px}
  .cover .meta{color:#666;font-size:13px;margin-top:30px;line-height:2}
  .footer{margin-top:50px;padding-top:15px;border-top:1px solid #ccc;
          text-align:center;color:#999;font-size:11px}
  @media print{body{margin:0;max-width:none} .no-print{display:none}}
</style>
"""


def _html_img(rel: str, width: str = "100%", caption: str = "") -> str:
    """嵌入图片为 base64，使 HTML 单文件可独立打开。"""
    import base64
    p = _img_path(rel)
    if not os.path.exists(p):
        return ""
    with open(p, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    ext = rel.split(".")[-1].lower()
    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg"}.get(ext, "image/png")
    cap = f'<div style="text-align:center;color:#888;font-size:12px;margin-bottom:10px">{caption}</div>' if caption else ""
    return f'<div style="text-align:center"><img src="data:{mime};base64,{b64}" style="width:{width};max-width:600px"></div>{cap}'


def _md_to_html_body(md_text: str) -> str:
    """简易 Markdown → HTML（不依赖外部库，覆盖报告用到的语法）。"""
    import re as _re
    lines = md_text.split("\n")
    html = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        # 表格
        if "|" in stripped and stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(_re.match(r"^[-: ]+$", c) for c in cells):
                continue  # 分隔行
            if not in_table:
                html.append("<table>")
                in_table = True
                html.append("<tr>" + "".join(f"<th>{c}</th>" for c in cells) + "</tr>")
                continue
            html.append("<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
            continue
        if in_table:
            html.append("</table>")
            in_table = False
        # 标题
        m = _re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            lvl = len(m.group(1))
            html.append(f"<h{lvl}>{m.group(2)}</h{lvl}>")
            continue
        # 分隔线
        if stripped == "---":
            html.append("<hr>")
            continue
        # 引用
        if stripped.startswith(">"):
            html.append(f"<blockquote>{stripped[1:].strip()}</blockquote>")
            continue
        # 列表
        if _re.match(r"^[-*]\s+", stripped):
            html.append(f"<li>{stripped[3:]}</li>")
            continue
        if stripped:
            html.append(f"<p>{stripped}</p>")
    if in_table:
        html.append("</table>")
    return "\n".join(html)


def generate_risk_html(rid: str) -> str:
    md = generate_risk_markdown(rid)
    body = _md_to_html_body(md)
    # 插入图片
    ctx = collect_risk_context(rid)
    geo = ctx["geophysics"]
    img_html = ""
    if geo:
        img_html = _html_img(geo["image"], "80%", geo["name"])
    for b in ctx["boreholes"]:
        img_html += _html_img(f"boreholes/{b['id']}.png", "45%", f"{b['id']} 钻孔柱状图")
    return f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">
<title>{ctx['risk']['name']} 风险分析报告</title>{HTML_CSS}</head>
<body>{body}{img_html}<div class="footer">
本报告由多源勘察数据联动展示与证据链追溯系统自动生成 · {_now()}</div>
</body></html>"""


def generate_full_html() -> str:
    md = generate_full_markdown()
    body = _md_to_html_body(md)
    # 插入所有物探与钻孔图
    img_html = '<h2>附图：物探剖面与钻孔柱状图</h2>'
    for g in GEO_LINES:
        img_html += _html_img(g["image"], "80%", g["name"])
    for b in BOREHOLES:
        img_html += _html_img(f"boreholes/{b['id']}.png", "40%", f"{b['id']}")
    return f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">
<title>{MANIFEST['project']['scenario']} 综合风险分析报告</title>{HTML_CSS}</head>
<body>{body}{img_html}<div class="footer">
本报告由多源勘察数据联动展示与证据链追溯系统自动生成 · {_now()}</div>
</body></html>"""


# #############################################################################
#  E. 统一入口
# #############################################################################
def generate_report(scope: str = "full", rid: Optional[str] = None,
                    fmt: str = "md") -> Dict[str, Any]:
    """统一报告生成入口。
    scope: 'full' 全线 | 'risk' 单风险（需提供 rid）
    fmt:   'md' | 'docx' | 'html'
    返回：{filename, content (bytes/str), media_type}
    """
    if scope == "risk":
        if not rid or rid not in RISK_BY_ID:
            raise ValueError(f"无效的风险 ID: {rid}")
        name = RISK_BY_ID[rid]["mileage"] + "_" + RISK_BY_ID[rid]["type_cn"]
        if fmt == "md":
            return {"filename": f"{name}_风险分析报告.md",
                    "content": generate_risk_markdown(rid).encode("utf-8-sig"),
                    "media_type": "text/markdown; charset=utf-8"}
        elif fmt == "docx":
            return {"filename": f"{name}_风险分析报告.docx",
                    "content": generate_risk_docx(rid),
                    "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        elif fmt == "html":
            return {"filename": f"{name}_风险分析报告.html",
                    "content": generate_risk_html(rid).encode("utf-8"),
                    "media_type": "text/html; charset=utf-8"}
    else:  # full
        if fmt == "md":
            return {"filename": "全线综合风险分析报告.md",
                    "content": generate_full_markdown().encode("utf-8-sig"),
                    "media_type": "text/markdown; charset=utf-8"}
        elif fmt == "docx":
            return {"filename": "全线综合风险分析报告.docx",
                    "content": generate_full_docx(),
                    "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        elif fmt == "html":
            return {"filename": "全线综合风险分析报告.html",
                    "content": generate_full_html().encode("utf-8"),
                    "media_type": "text/html; charset=utf-8"}
    raise ValueError(f"不支持的格式: {fmt}")


# #############################################################################
#  F. 预览（用于前端展示，不下载）
# #############################################################################
def preview_report(scope: str = "full", rid: Optional[str] = None) -> Dict[str, Any]:
    """返回报告的 Markdown 与 HTML 预览，供前端在线查看。"""
    if scope == "risk":
        if not rid or rid not in RISK_BY_ID:
            raise ValueError(f"无效的风险 ID: {rid}")
        md = generate_risk_markdown(rid)
    else:
        md = generate_full_markdown()
    html_body = _md_to_html_body(md)
    return {"markdown": md, "html": html_body,
            "title": (RISK_BY_ID[rid]["name"] if scope == "risk" and rid in RISK_BY_ID
                      else MANIFEST["project"]["scenario"] + " 综合报告")}


# ----------------------------------------------------------------------------
# 自测
# ----------------------------------------------------------------------------
if __name__ == "__main__":
    print("=" * 60)
    print("报告生成引擎自测")
    print("=" * 60)
    # Markdown
    md1 = generate_risk_markdown("R001")
    md2 = generate_full_markdown()
    print(f"[md] 单风险 R001: {len(md1)} 字符")
    print(f"[md] 全线综合: {len(md2)} 字符")
    # DOCX
    d1 = generate_risk_docx("R001")
    d2 = generate_full_docx()
    print(f"[docx] 单风险 R001: {len(d1)//1024} KB")
    print(f"[docx] 全线综合: {len(d2)//1024} KB")
    # HTML
    h1 = generate_risk_html("R001")
    h2 = generate_full_html()
    print(f"[html] 单风险 R001: {len(h1)//1024} KB")
    print(f"[html] 全线综合: {len(h2)//1024} KB")
    # 统一入口
    r = generate_report("full", fmt="docx")
    print(f"[入口] full/docx -> {r['filename']} ({len(r['content'])//1024} KB)")
    r = generate_report("risk", rid="R002", fmt="md")
    print(f"[入口] risk/R002/md -> {r['filename']}")
    print("\n✓ 报告生成引擎自测通过")
